import codecs
from flask import Flask, send_from_directory, request, json
import logging
import os
from flask_cors import CORS, cross_origin
import cv2
import numpy as np
import os.path
import io
# import re
from docx   import Document
from pytesseract import pytesseract
from werkzeug.utils import secure_filename
import enchant

app = Flask(__name__)
CORS(app)
file_handler = logging.FileHandler('server.log')
app.logger.addHandler(file_handler)
app.logger.setLevel(logging.INFO)

PROJECT_HOME = os.path.dirname(os.path.realpath(__file__))
UPLOAD_FOLDER = '{}/uploads/'.format(PROJECT_HOME)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


def create_new_folder(local_dir):
    newpath = local_dir
    if not os.path.exists(newpath):
        os.makedirs(newpath)
    return newpath

@app.route('/process_image', methods = ['POST', 'GET'])
def api_process_image():
    app.logger.info(PROJECT_HOME)
    if request.method == 'POST' and request.files['image']:
        app.logger.info(app.config['UPLOAD_FOLDER'])
        img = request.files['image']
        img_name = request.form['name_image']
        in_memory_file = io.BytesIO()
        img.save(in_memory_file)
        data = np.fromstring(in_memory_file.getvalue(), dtype=np.uint8)
        color_image_flag = 1
        img_opencv = cv2.imdecode(data, color_image_flag)
        us_dict = enchant.Dict("en_US")
        result_text = convert_imge_text(img_opencv)
        img_name = secure_filename(img_name)
        create_new_folder(app.config['UPLOAD_FOLDER'])
        document = Document()
        document.add_heading('thach created this', 0)
        # print(result_text)
        result_array = result_text.split('\n')
        for textLine in result_array:
            p = document.add_paragraph('')
            word_array = textLine.split(' ')
            for word in word_array:
                if not word:
                    continue
                if us_dict.check(word):
                    p.add_run(word)
                    p.add_run(' ')
                else:
                    p.add_run(word).bold = True
                    p.add_run(' ')
            # document.add_paragraph(textLine)
        saved_path = os.path.join(app.config['UPLOAD_FOLDER'], img_name.split('.')[0] + ".docx")
        document.save(saved_path)
        app.logger.info("saving {}".format(saved_path))
        return json.dumps({'success': True, 'filename': img_name.split('.')[0] + '.docx'}), 200, {'ContentType': 'application/json'}
    else:
        return json.dumps({'success': False}), 400, {'ContentType': 'application/json'}

@app.route('/download', methods = ['POST', 'GET'])
def api_root():
    docx_name = request.args.get('docx_name')
    print(docx_name)
    return send_from_directory(app.config['UPLOAD_FOLDER'], docx_name, as_attachment=True)

def convert_imge_text(input_img):
    DEBUG = 0


    # pixel intensity = 0.30R + 0.59G + 0.11B
    def ii(xx, yy):
        if yy >= img_y or xx >= img_x:
            return 0
        pixel = img[yy][xx]
        return 0.30 * pixel[2] + 0.59 * pixel[1] + 0.11 * pixel[0]

    def connected(contour):
        first = contour[0][0]
        last = contour[len(contour) - 1][0]
        return abs(first[0] - last[0]) <= 1 and abs(first[1] - last[1]) <= 1

    def c(index):
        return contours[index]


    def count_children(index, h_, contour):
        if h_[index][2] < 0:
            return 0
        else:
            if keep(c(h_[index][2])):
                count = 1
            else:
                count = 0
            count += count_siblings(h_[index][2], h_, contour, True)
            return count


    def is_child(index, h_):
        return get_parent(index, h_) > 0


    def get_parent(index, h_):
        parent = h_[index][3]
        while not keep(c(parent)) and parent > 0:
            parent = h_[parent][3]

        return parent


    def count_siblings(index, h_, contour, inc_children=False):
        if inc_children:
            count = count_children(index, h_, contour)
        else:
            count = 0
        p_ = h_[index][0]
        while p_ > 0:
            if keep(c(p_)):
                count += 1
            if inc_children:
                count += count_children(p_, h_, contour)
            p_ = h_[p_][0]
        n = h_[index][1]
        while n > 0:
            if keep(c(n)):
                count += 1
            if inc_children:
                count += count_children(n, h_, contour)
            n = h_[n][1]
        return count


    def keep(contour):
        return keep_box(contour) and connected(contour)


    def keep_box(contour):
        xx, yy, w_, h_ = cv2.boundingRect(contour)

        w_ *= 1.0
        h_ *= 1.0

        if w_ / h_ < 0.1 or w_ / h_ > 10:
            if DEBUG:
                print("\t debug=> shape is invalid: (" + str(xx) + "," + str(yy) + "," + str(w_) + "," + str(h_) + ")" + \
                      str(w_ / h_))
            return False

        if ((w_ * h_) > ((img_x * img_y) / 5)) or ((w_ * h_) < 15):
            if DEBUG:
                print(w_, h_)
                print("\t debug=> size is invalid")
            return False

        return True


    def include_box(index, h_, contour):
        if DEBUG == 0 and index == 109:
            print(str(index) + ":")
            print(count_children(index, h_, contour))
            if is_child(index, h_):
                print("\tIs a child")
                print("\tparent " + str(get_parent(index, h_)) + " has " + str(
                    count_children(get_parent(index, h_), h_, contour)) + " children")
                print("\thas " + str(count_children(index, h_, contour)) + " children")

        if is_child(index, h_) and count_children(get_parent(index, h_), h_, contour) <= 2:
            if DEBUG:
                print("\t debug: is an interior to a letter")
            return False

        if count_children(index, h_, contour) > 2:
            if not DEBUG:
                print(str(index) + ":")
                print("\t debug: is a container of letters")
            return False

        if DEBUG:
            print("\t debug: keeping")
        return True

    output_file = "output.png"
    kernel = np.ones((1, 1), np.uint8)
    # dilate_img = cv2.dilate(input_img, kernel, iterations=1)
    # final_img = cv2.erode(dilate_img, kernel, iterations=1)
    # increase_img = cv2.resize(input_img, None, fx=2.5, fy=2.5, interpolation=cv2.INTER_LINEAR)
    dilate_img = cv2.dilate(input_img, kernel, iterations=13)
    final_img = cv2.erode(dilate_img, kernel, iterations=1)

    img = cv2.copyMakeBorder(final_img, 50, 50, 50, 50, cv2.BORDER_CONSTANT)

    img_y = len(img)
    img_x = len(img[0])

    if DEBUG:
        print("Image is " + str(len(img)) + "x" + str(len(img[0])))

    blue, green, red = cv2.split(img)

    blue_edges = cv2.Canny(blue, 200, 250)
    green_edges = cv2.Canny(green, 200, 250)
    red_edges = cv2.Canny(red, 200, 250)

    edges = blue_edges | green_edges | red_edges

    image, contours, hierarchy = cv2.findContours(edges.copy(), cv2.RETR_TREE, cv2.CHAIN_APPROX_NONE)

    hierarchy = hierarchy[0]
    print(hierarchy)
    if DEBUG:
        processed = edges.copy()
        rejected = edges.copy()

    keepers = []
    # print(contours)
    for index_, contour_ in enumerate(contours):
        if DEBUG:
            print("Processing #%d" % index_)

        x, y, w, h = cv2.boundingRect(contour_)

        # Check the contour and it's bounding box
        if keep(contour_) and include_box(index_, hierarchy, contour_):
            # It's a winner!
            keepers.append([contour_, [x, y, w, h]])
            if DEBUG:
                cv2.rectangle(processed, (x, y), (x + w, y + h), (100, 100, 100), 1)
                cv2.putText(processed, str(index_), (x, y - 5), cv2.FONT_HERSHEY_PLAIN, 1, (255, 255, 255))
        else:
            if DEBUG:
                cv2.rectangle(rejected, (x, y), (x + w, y + h), (100, 100, 100), 1)
                cv2.putText(rejected, str(index_), (x, y - 5), cv2.FONT_HERSHEY_PLAIN, 1, (255, 255, 255))

    new_image = edges.copy()
    new_image.fill(255)
    boxes = []

    for index_, (contour_, box) in enumerate(keepers):

        fg_int = 0.0
        for p in contour_:
            fg_int += ii(p[0][0], p[0][1])

        fg_int /= len(contour_)
        if DEBUG:
            print("FG Intensity for #%d = %d" % (index_, fg_int))

        x_, y_, width, height = box
        bg_int = \
            [
                # bottom left corner 3 pixels
                ii(x_ - 1, y_ - 1),
                ii(x_ - 1, y_),
                ii(x_, y_ - 1),

                # bottom right corner 3 pixels
                ii(x_ + width + 1, y_ - 1),
                ii(x_ + width, y_ - 1),
                ii(x_ + width + 1, y_),

                # top left corner 3 pixels
                ii(x_ - 1, y_ + height + 1),
                ii(x_ - 1, y_ + height),
                ii(x_, y_ + height + 1),

                # top right corner 3 pixels
                ii(x_ + width + 1, y_ + height + 1),
                ii(x_ + width, y_ + height + 1),
                ii(x_ + width + 1, y_ + height)
            ]
        for i in range(x_ + 1, x_ + width - 1):
            # print(i)
            bg_int.append(ii(i, y_ - 1))
            bg_int.append(ii(i, y_ + height + 1))
        bg_int = np.median(bg_int)

        if DEBUG:
            print("BG Intensity for #%d = %s" % (index_, repr(bg_int)))
        if fg_int >= bg_int:
            fg = 255
            bg = 0
        else:
            fg = 0
            bg = 255
        for x in range(x_, x_ + width):
            for y in range(y_, y_ + height):
                if y >= img_y or x >= img_x:
                    if DEBUG:
                        print("pixel out of bounds (%d,%d)" % (y, x))
                    continue
                if ii(x, y) > fg_int:
                    new_image[y][x] = bg
                else:
                    new_image[y][x] = fg
    new_image = cv2.blur(new_image, (2, 2))
    cv2.imwrite(output_file, new_image)
    result = pytesseract.image_to_string(new_image)
    file = codecs.open("output-test.txt", "w", "utf-8")
    file.write(result)
    # print(repr(result))
    if DEBUG:
        cv2.imwrite('edges.png', edges)
        cv2.imwrite('processed.png', processed)
        cv2.imwrite('rejected.png', rejected)
    return result
if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=False)
